"""向 docs/协作留痕.docx 追加本次对话留痕.

本次内容: 对话 70 · 2026-08-28 · 批量创建 7 个 GitLab 用户并分配 Developer 角色
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

DOC = Path(__file__).resolve().parents[1] / "docs" / "协作留痕.docx"


def main() -> None:
    doc = Document(DOC)

    last = 0
    for p in doc.paragraphs:
        m = re.search(r"对话\s*(\d+)", p.text)
        if m:
            last = max(last, int(m.group(1)))
    dialog_no = last + 1
    date = "2026-08-28"

    def add(text: str, bold: bool = False, color: tuple[int, int, int] | None = None,
            size: int = 11) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = RGBColor(*color)

    add(f"【用户】 批量创建 7 个 GitLab 用户并分配 Developer 角色 · 对话 {dialog_no} · {date}",
        bold=True, color=(0x1F, 0x49, 0x7D))
    add(
        "用户诉求: 在本地 GitLab 创建 7 个账号 (xuhao/lishuaishuai/gaoqianqian/lichuanhao/"
        "xiaojuan/sunzhouqing/zhuyushuang), 对应邮箱 (xuhao@jtexpress.com 等, 部分用英文名 "
        "kenna.gao/lennard.li/joy.sun), 初始密码 AutoDev2026!Xq, 分配项目 Developer 角色."
    )
    add("文件变更表:")
    add("  - scripts/create_gitlab_users.py (新增, 批量创建用户脚本, 支持 --dry-run)")
    add("  - docs/协作留痕.docx (本条追加)")
    add("验证结果表:")
    add("  - PAT 验证: curl /api/v4/user 返回 200, root PAT 仍可用")
    add("  - dry-run: 7 用户全部预演成功")
    add("  - 实际创建: 7/7 用户创建成功 (id=6-12), 7/7 加为项目 Developer 成功")
    add("  - 验证成员列表: /projects/1/members 返回 9 人 (root Owner + team + 7 新建 Developer)")
    add("影响范围分析表:")
    add("  - 7 个同事可用账号 + AutoDev2026!Xq 登录 http://10.66.67.26:8929")
    add("  - 拥有 root/automation-platform 项目 Developer 权限 (可推 master, 不能改 protected branches)")
    add("  - 同事登录后可在 http://10.66.67.26:8080 跑测试 (前端工作台, 需先放行 8080 防火墙)")

    add(f"【助手】 用 GitLab API 批量创建 · 对话 {dialog_no} · {date}",
        bold=True, color=(0x1F, 0x49, 0x7D))
    add(
        "方案: 用 root 的 PAT (glpat-y0Q5yhq9iFvWxUjoj7ZM6286MQp1OjEH.01.0w0nsh5or) "
        "调 GitLab REST API. 两步: (1) POST /users 创建用户 (skip_confirmation 跳过邮箱确认, "
        "reset_password=False 不强制改密); (2) POST /projects/1/members 加为项目 Developer "
        "(access_level=30). 写 scripts/create_gitlab_users.py, 用 urllib 不依赖 requests."
    )
    add(
        "脚本设计: 用户列表 USERS 内置 7 个 (username/name/email 三元组); "
        "--dry-run 选项先预演; 已存在用户 (409) 跳过; 已是成员 (409) 跳过; "
        "GITLAB_TOKEN / GITLAB_URL 支持环境变量覆盖. "
        "初始密码 AutoDev2026!Xq, skip_confirmation=True 跳过邮箱确认, "
        "reset_password=False 不强制首次登录改密 (用户可后续在 Preferences 改)."
    )
    add(
        "GitLab access_level 对照: 10 Guest / 20 Reporter / 30 Developer / 40 Maintainer / 50 Owner. "
        "选 30 Developer: 同事能 clone/push master 分支、提 MR, 但不能改 protected branches、不能加成员、不能删项目. "
        "这是 '部门同事协同开发' 的合理权限级别."
    )
    add(
        "验证: GET /projects/1/members 返回 9 人, root (Owner) + team (Developer, 之前创建的共享账号) "
        "+ 7 新建用户 (Developer). 角色分配正确."
    )
    add(
        "未做但可考虑的事: (1) 强制首次登录改密 (reset_password=True, 更安全但同事要改一次); "
        "(2) 给同事发邮件通知账号 (需配 SMTP); (3) 同事登录后默认分支看到 master (GitLab 默认分支已是 master); "
        "(4) 同事用 SSH 推送需各自上传公钥到 User Settings → SSH Keys."
    )
    add("文件变更表:")
    add("  - scripts/create_gitlab_users.py (新增, 120 行 Python, 用 urllib 调 GitLab API)")
    add("  - docs/协作留痕.docx (本条追加)")
    add("  - GitLab 服务器: 新增 7 个用户 (id=6-12), 7 个 Developer 成员关系")
    add("验证结果表:")
    add("  - scripts/create_gitlab_users.py --dry-run 输出 7 个用户预演成功")
    add("  - scripts/create_gitlab_users.py 实际跑: 7/7 创建 + 7/7 加成员成功")
    add("  - GET /projects/1/members 验证 9 人成员列表, 角色分配正确")
    add("影响范围分析表:")
    add("  - 7 个同事: 用 账号 / AutoDev2026!Xq 登录 http://10.66.67.26:8929 即可使用")
    add("  - 同事权限: Developer, 可 clone/push master, 提 MR, 但不能改项目设置")
    add("  - 安全: 初始密码统一, 建议同事首次登录后改密 (Preferences → Password)")
    add("  - 待办: 用户需放行 Windows 防火墙 8080 端口 (管理员 PowerShell) 才能让同事访问前端工作台")

    doc.save(DOC)
    print(f"已追加对话 {dialog_no} 到 {DOC}")
    print(f"文件大小: {DOC.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
