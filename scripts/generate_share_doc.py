"""生成《自动化测试平台分享文档.docx》.

面向部门内测试人员(部分非开发出身)的分享文档, 介绍从零搭建的全过程、
框架设计思路、技术栈选型、代码管理、部署方案与同事参与方式.

运行:
    .venv/Scripts/python.exe scripts/generate_share_doc.py
产物:
    docs/自动化测试平台分享文档.docx
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "自动化测试平台分享文档.docx"

# ---------- 样式工具 ----------


def set_cell_bg(cell: Any, color_hex: str) -> None:
    """设置单元格背景色."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_run_font(run: Any, name: str = "微软雅黑", size: int = 11, bold: bool = False,
                 color: RGBColor | None = None) -> None:
    """统一设置字体: 中文微软雅黑, 西文同体."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def add_heading(doc: Document, text: str, level: int = 1) -> Any:
    """添加标题, 统一样式."""
    sizes = {0: 26, 1: 20, 2: 16, 3: 14, 4: 12}
    colors = {0: RGBColor(0x1F, 0x49, 0x7D), 1: RGBColor(0x1F, 0x49, 0x7D),
              2: RGBColor(0x2E, 0x74, 0xB5), 3: RGBColor(0x2E, 0x74, 0xB5),
              4: RGBColor(0x40, 0x40, 0x40)}
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 0 else WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, name="微软雅黑", size=sizes.get(level, 12), bold=True,
                 color=colors.get(level))
    return p


def add_para(doc: Document, text: str, size: int = 11, bold: bool = False,
             indent: float = 0.0, color: RGBColor | None = None) -> Any:
    """添加段落正文."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    if indent > 0:
        p.paragraph_format.first_line_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_code_block(doc: Document, code: str) -> Any:
    """添加代码块: 灰色背景等宽字体."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.2
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(code)
    set_run_font(run, name="Consolas", size=10, color=RGBColor(0x33, 0x33, 0x33))
    return p


def add_bullet(doc: Document, text: str, level: int = 0) -> Any:
    """添加无序列表项."""
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              col_widths: list[float] | None = None) -> Any:
    """添加表格: 表头深蓝底白字."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    table.autofit = True

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, "1F497D")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(val)
            set_run_font(run, size=10)
            if r_idx % 2 == 0:
                set_cell_bg(cell, "F2F2F2")

    return table


def add_callout(doc: Document, text: str, label: str = "提示") -> Any:
    """添加提示框: 浅黄背景."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, "FFF8E1")
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.5
    run_label = p.add_run(f"【{label}】 ")
    set_run_font(run_label, size=11, bold=True, color=RGBColor(0xE6, 0x7E, 0x22))
    run = p.add_run(text)
    set_run_font(run, size=11)
    # 加点空白行
    doc.add_paragraph()
    return table


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------- 文档结构 ----------


def build_cover(doc: Document) -> None:
    """封面."""
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("自动化测试平台")
    set_run_font(run, size=36, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("从零搭建全历程分享")
    set_run_font(run, size=24, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— 部门内部技术分享 ——")
    set_run_font(run, size=14, color=RGBColor(0x80, 0x80, 0x80))

    for _ in range(8):
        doc.add_paragraph()

    for label, value in [("分享人", "wsy"), ("日 期", "2026 年 8 月"),
                          ("受 众", "部门测试人员(含非开发背景同事)"),
                          ("文档版本", "V1.0")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}：{value}")
        set_run_font(run, size=12, color=RGBColor(0x40, 0x40, 0x40))


def build_toc(doc: Document) -> None:
    """目录说明(静态文本, 不依赖 Word 域)."""
    add_heading(doc, "目  录", level=0)
    doc.add_paragraph()
    toc_items = [
        "第一部分  破题: 为什么要搭建这个平台",
        "   1.1  之前的痛点",
        "   1.2  目标与定位",
        "   1.3  谁会用, 怎么用",
        "第二部分  总览: 整体架构与技术栈",
        "   2.1  分层架构图",
        "   2.2  技术栈选型一览表",
        "   2.3  为什么选这些技术",
        "第三部分  从零搭建的全过程(14 个阶段)",
        "   3.1  阶段 0-1: 打地基, 立骨架",
        "   3.2  阶段 2-4: 三大基础能力",
        "   3.3  阶段 5-8: 报告/扩展/CI/AI",
        "   3.4  阶段 9-13: 收尾与平台化",
        "第四部分  核心设计思路详解",
        "   4.1  framework + projects 的分层逻辑",
        "   4.2  数据驱动与接口依赖",
        "   4.3  并发执行与 Allure 报告",
        "   4.4  WebUI/App 扩展",
        "   4.5  AI 失败分析",
        "第五部分  一个贯穿示例: 从写用例到看报告",
        "   5.1  写一个登录接口测试",
        "   5.2  跑起来",
        "   5.3  看 Allure 报告",
        "第六部分  代码管理与协同",
        "   6.1  三套 Git 远程: 个人 / 本地 GitLab",
        "   6.2  分支与提交规范",
        "   6.3  Pre-commit 自动检查",
        "   6.4  CI/CD 流水线",
        "第七部分  部署与运维",
        "   7.1  本地依赖(MySQL/Redis/Kafka)",
        "   7.2  测试管理平台一键启动",
        "   7.3  本地 GitLab 容器",
        "   7.4  数据备份与镜像管理",
        "第八部分  同事参与指南",
        "   8.1  新人 30 分钟上手",
        "   8.2  跑测试的三种姿势",
        "   8.3  查看报告",
        "   8.4  提交代码流程",
        "第九部分  踩过的坎与决策理由",
        "第十部分  后续演进与展望",
        "附录 A  技术名词解释",
        "附录 B  常用命令速查",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.4
        run = p.add_run(item)
        set_run_font(run, size=11,
                     color=RGBColor(0x33, 0x33, 0x33) if item.startswith(" ")
                     else RGBColor(0x1F, 0x49, 0x7D),
                     bold=not item.startswith(" "))
    add_page_break(doc)


# ---------- 各章节 ----------


def build_part1(doc: Document) -> None:
    add_heading(doc, "第一部分  破题: 为什么要搭建这个平台", level=1)

    add_heading(doc, "1.1  之前的痛点", level=2)
    add_para(doc,
        "在我们做自动化测试之前, 部门里有几个长期困扰的问题, 这也是我决定搭建这个平台的初衷:",
        indent=0.7)
    add_bullet(doc, "每个业务系统(OMS/WMS/TMS/JMS 系列)各写一套测试脚本, 能力重复造轮子")
    add_bullet(doc, "标准不统一: 这个人用 requests, 那个人用 httpx; 这个人用 unittest, 那个人用 pytest")
    add_bullet(doc, "测试数据散落, 环境配置硬编码在代码里, 切换环境要改源码")
    add_bullet(doc, "无统一报告: 测试结果只能看终端输出, 没有截图、步骤、历史趋势")
    add_bullet(doc, "CI 接入成本高: 每个项目都要重新写流水线脚本")
    add_bullet(doc, "新人上手慢: 进项目第一周基本在搞环境, 没法立即出活")

    add_heading(doc, "1.2  目标与定位", level=2)
    add_para(doc, "一句话定位: 一个企业级 Python 自动化测试平台, 共享框架内核, 各业务系统独立维护测试代码.",
             indent=0.7, bold=True)

    add_para(doc, "具体目标拆成 5 条:", indent=0.7)
    add_bullet(doc, "共享框架内核(framework), 各业务系统独立维护测试代码(projects)")
    add_bullet(doc, "覆盖 API / WebUI / App 三种自动化类型(API 优先)")
    add_bullet(doc, "支持 DB / Redis / Kafka / RabbitMQ / RocketMQ 数据校验")
    add_bullet(doc, "支持数据驱动、接口依赖编排、并发执行、Allure 报告")
    add_bullet(doc, "支持 GitLab CI / Jenkins / Docker / 本地 GitLab 协同")

    add_callout(doc,
        "本平台是测试工作台, 不是 CI 调度中心. 它解决的是\"怎么方便地跑测试、看结果\", "
        "定时调度、远程分发、权限控制属于后续演进, 不在第一版范围内.",
        label="定位边界")

    add_heading(doc, "1.3  谁会用, 怎么用", level=2)
    add_table(doc,
        headers=["角色", "使用方式", "需要的技能"],
        rows=[
            ["测试开发", "改 framework / 加业务测试工程", "Python 进阶, 框架设计"],
            ["测试工程师", "写用例、跑测试、看报告", "Python 基础, pytest 用法"],
            ["测试人员(非开发)", "用 Web 平台触发测试、查看报告", "浏览器操作即可, 无需写代码"],
            ["运维/CI 维护", "维护 CI 流水线与部署", "Docker, GitLab CI/Jenkins"],
        ],
        col_widths=[3.5, 6.5, 5.0])
    add_para(doc, "本文档面向所有人, 重点照顾非开发出身的测试同事, 会把技术名词都解释清楚.",
             indent=0.7)
    add_page_break(doc)


def build_part2(doc: Document) -> None:
    add_heading(doc, "第二部分  总览: 整体架构与技术栈", level=1)

    add_heading(doc, "2.1  分层架构图", level=2)
    add_para(doc,
        "整个平台分三层: 业务测试工程层(projects)、共享框架层(framework)、基础设施层(CI/Docker/工具链).",
        indent=0.7)
    add_code_block(doc, """┌─────────────────────────────────────────────────────┐
│              projects/ (业务测试工程)                 │
│  jmseg(API/埃及)  jmseu(WebUI/德国)  oms  wms  tms  │
│  每个: api/ pages/ testcase/ data/ fixture/ config/ │
├─────────────────────────────────────────────────────┤
│              framework/ (共享框架核心)                │
│  ┌──────────┬──────────┬──────────┬──────────────┐ │
│  │  core/   │ clients/ │ testing/ │  reporting/  │ │
│  │ 配置日志 │ HTTP/Web │ 数据驱动 │  Allure扩展  │ │
│  │ 上下文   │ DB/Cache │ 接口依赖 │  标签/环境   │ │
│  │ 异常     │ MQ/App   │ 断言提取 │  通知/历史   │ │
│  ├──────────┴──────────┴──────────┴──────────────┤ │
│  │              plugins/  utils/                   │ │
│  │           AI失败分析  通用工具                   │ │
│  └─────────────────────────────────────────────────┘ │
├─────────────────────────────────────────────────────┤
│         基础设施 (CI/CD / Docker / 工具链)            │
│  GitLab CI  Jenkins  Docker  uv  ruff  mypy  pre-commit │
└─────────────────────────────────────────────────────┘""")

    add_callout(doc,
        "最关键的分层原则: framework 不依赖 projects, 下层不依赖上层. "
        "业务测试代码不能反向污染共享框架, 这样框架能力可以稳步沉淀, 不被业务绑死.",
        label="核心原则")

    add_heading(doc, "2.2  技术栈选型一览表", level=2)
    add_table(doc,
        headers=["类别", "技术", "版本", "选型理由"],
        rows=[
            ["语言", "Python", ">=3.12", "asyncio 原生支持、类型系统成熟"],
            ["测试框架", "pytest", ">=8.3", "生态最丰富、fixture 机制灵活"],
            ["依赖管理", "uv", "最新", "比 pip/poetry 快 10-100 倍、workspace 原生支持"],
            ["HTTP 客户端", "httpx", ">=0.27", "原生 async、HTTP/2、mock transport 支持好"],
            ["Web 自动化", "Playwright", ">=1.40", "跨浏览器、原生 async、自动等待、trace viewer"],
            ["App 自动化", "Appium", ">=2.0", "跨平台、社区成熟"],
            ["配置", "pydantic v2 + pydantic-settings", ">=2.7", "类型安全、多源合并、env 覆盖"],
            ["日志", "loguru", ">=0.7", "零配置、结构化、文件轮转"],
            ["数据库", "SQLAlchemy", ">=2.0", "async engine、多方言、连接池"],
            ["缓存", "redis", ">=5.0", "async 支持、连接池"],
            ["消息队列", "aiokafka / aio-pika / rocketmq", "最新", "三端 async 客户端"],
            ["报告", "allure-pytest", ">=2.13", "行业标准、富文本报告、历史趋势"],
            ["并发", "pytest-xdist", ">=3.5", "进程级并行、隔离性好"],
            ["代码质量", "ruff + mypy", "最新", "ruff 替代 flake8+isort+black, mypy strict"],
            ["CI/CD", "GitLab CI + Jenkins", "-", "双流水线共用 scripts/ci/ci.sh"],
            ["平台后端", "FastAPI", ">=0.110", "原生 async、自动生成 OpenAPI 文档"],
            ["平台前端", "Vue 3 + Element Plus + Vite", "最新", "组合式 API、构建快、组件齐全"],
            ["平台元数据库", "MySQL", "8.4", "持久化项目、用例、执行记录、报告快照"],
            ["容器化", "Docker / docker-compose", "最新", "本地依赖编排 + GitLab 部署"],
        ],
        col_widths=[2.5, 4.5, 2.0, 6.0])

    add_heading(doc, "2.3  为什么选这些技术", level=2)
    add_para(doc,
        "技术选型不是追新, 每个都有具体的选型理由. 这里挑几个非开发同事可能看不太懂的关键点展开讲:",
        indent=0.7)

    add_para(doc, "(1) 为什么用 uv 不用 pip?", bold=True, indent=0.7)
    add_para(doc,
        "uv 是 Astral 公司用 Rust 写的 Python 包管理器, 比 pip 快 10-100 倍. "
        "另外它原生支持 workspace, 让 framework 和 projects 共享一个虚拟环境, "
        "不用每个项目都装一遍依赖, 这是 pip 做不到的.", indent=0.7)

    add_para(doc, "(2) 为什么用 pytest 不用 unittest?", bold=True, indent=0.7)
    add_para(doc,
        "pytest 的 fixture 机制可以让测试前置/后置操作非常灵活, 而且生态庞大, "
        "可以接 allure 报告、pytest-xdist 并发、参数化等. unittest 是 Python 自带的, "
        "功能够用但扩展性差.", indent=0.7)

    add_para(doc, "(3) 为什么用 httpx 不用 requests?", bold=True, indent=0.7)
    add_para(doc,
        "requests 是同步库, httpx 原生支持 async(异步). 我们的测试会并发跑很多请求, "
        "异步能显著提速. 另外httpx自带 MockTransport, 测试时可以不连真实后端, "
        "这对 CI 环境非常重要.", indent=0.7)

    add_para(doc, "(4) 为什么用 Playwright 不用 Selenium?", bold=True, indent=0.7)
    add_para(doc,
        "Playwright 微软出品, 跨浏览器、原生 async、自动等待(不用手写 sleep)、"
        "trace viewer 调试神器. Selenium 老牌但 API 设计偏旧, 异步支持差.", indent=0.7)

    add_para(doc, "(5) 为什么 framework + projects 分层?", bold=True, indent=0.7)
    add_para(doc,
        "如果不分层, 每个业务系统都自己写一套 HTTP 客户端、日志、断言, 重复造轮子. "
        "分层后 framework 沉淀公共能力, projects 只写自己的业务逻辑, "
        "新人进项目不用重学, 框架升级一处生效所有项目受益.", indent=0.7)
    add_page_break(doc)


def build_part3(doc: Document) -> None:
    add_heading(doc, "第三部分  从零搭建的全过程(14 个阶段)", level=1)

    add_para(doc,
        "这个平台不是一蹴而就的, 一共拆成 14 个阶段逐步演进. 这一节讲每个阶段做了什么、"
        "遇到什么问题、为什么这样决策. 阶段顺序就是搭建顺序, 也是学习顺序.",
        indent=0.7)

    add_heading(doc, "3.1  阶段 0-1: 打地基, 立骨架", level=2)
    add_table(doc,
        headers=["阶段", "内容", "关键产出"],
        rows=[
            ["0", "工程地基(uv / 工具链 / CI 骨架 / 冒烟测试)", "pyproject.toml, uv.lock, ruff/mypy 配置"],
            ["0+", "完整目录结构与基础设施文件", "framework/projects/tests 目录骨架"],
            ["1", "框架核心(config / logger / context / exceptions / registry)", "core 模块完成"],
        ],
        col_widths=[1.5, 7.0, 6.5])

    add_para(doc, "踩过的坎:", bold=True, indent=0.7)
    add_bullet(doc,
        "uv workspace 配置: 一开始不知道 framework 要在 [tool.uv.sources] 里写 "
        "{ workspace = true }, 导致 projects 找不到 framework 包")
    add_bullet(doc,
        "Python 3.12 + setuptools: PyCharm 测试运行器需要 setuptools<81, "
        "因为 v83 移除了 pkg_resources, 这个坑埋了很久才定位")
    add_bullet(doc, "mypy strict 模式: 一开始开启 strict 到处报错, 后来用 overrides 排除可选依赖")

    add_heading(doc, "3.2  阶段 2-4: 三大基础能力", level=2)
    add_table(doc,
        headers=["阶段", "内容", "关键产出"],
        rows=[
            ["2", "API 能力 + 基础设施(http client / extractors / assertions)",
             "AsyncHttpClient, TokenManager, BearerAuth"],
            ["3", "数据驱动 + 接口依赖(datadriven / dependency DAG / variables)",
             "load_cases, 拓扑排序, 接口依赖编排"],
            ["4", "DB / Cache / MQ(Kafka + RabbitMQ + RocketMQ 三端)",
             "DatabaseClient, CacheClient, MessageClient"],
        ],
        col_widths=[1.5, 7.0, 6.5])

    add_para(doc, "关键设计决策:", bold=True, indent=0.7)
    add_bullet(doc,
        "HTTP 客户端用 async context manager, 用完即关, 避免连接泄漏")
    add_bullet(doc,
        "数据驱动用 YAML 不用 Excel: YAML 易读、可嵌套、Git diff 友好; "
        "Excel 是二进制, 改一行无法 review")
    add_bullet(doc,
        "接口依赖用 DAG(有向无环图) + 拓扑排序: 登录拿 token → 创建订单 → 查询, "
        "自动按依赖顺序跑, 不用手动排顺序")
    add_bullet(doc,
        "MQ 三端用统一 MessageClient 接口: 业务代码不感知是 Kafka 还是 RabbitMQ, "
        "切消息中间件只改配置")

    add_heading(doc, "3.3  阶段 5-8: 报告/扩展/CI/AI", level=2)
    add_table(doc,
        headers=["阶段", "内容", "关键产出"],
        rows=[
            ["5", "并发执行(pytest-xdist) + Allure 报告(attach/labels/environment/categories)",
             "-n auto, allure step, 失败截图"],
            ["6", "Web/App 扩展(Playwright WebClient + Appium AppClient)",
             "含腾讯滑块验证码自动求解器"],
            ["7", "CI/CD + Docker(GitLab CI / Jenkins / Dockerfile / docker-compose)",
             "ci.sh 单入口, 双流水线共用"],
            ["8", "AI 失败分析(FailureAnalyzer ABC + NullAnalyzer + LLMAnalyzer)",
             "用例失败自动调 LLM 分析原因"],
        ],
        col_widths=[1.5, 7.0, 6.5])

    add_para(doc, "亮点:", bold=True, indent=0.7)
    add_bullet(doc,
        "Allure 报告自动附件: 测试跑完, HTTP 请求/响应(已脱敏)自动作为附件, 不用手写")
    add_bullet(doc,
        "腾讯滑块验证码: 用 OpenCV 缺口检测 + 人类轨迹模拟自动求解, "
        "但太慢且易被封 IP, 后来改成人机协作(操作员手动完成滑动, 脚本等待成功标志)")
    add_bullet(doc,
        "CI 单入口: scripts/ci/ci.sh 是唯一的命令入口, GitLab CI 和 Jenkins 都调用它, "
        "改命令只动一处")
    add_bullet(doc,
        "AI 失败分析: 用例失败后, LLMAnalyzer 把失败堆栈和日志喂给大模型, 自动给出可能原因. "
        "这个能力降低了排错门槛")

    add_heading(doc, "3.4  阶段 9-13: 收尾与平台化", level=2)
    add_table(doc,
        headers=["阶段", "内容", "关键产出"],
        rows=[
            ["9", "集成验证 + Allure 报告增强 + 文档同步", "完整文档体系"],
            ["10", "数据生命周期 + AI 用例生成 + 邮件通知/Allure 历史 + 业务系统脚手架(oms/wms/tms)",
             "数据清理, 历史趋势"],
            ["11", "接口压测引擎(Locust + YAML 场景 + 数据驱动 + SLA + HTML 报告)",
             "loadtest 项目, 阶梯式加压"],
            ["12", "数据源拉取脚本(生产运单库分页查询 -> CSV/TXT)", "生产数据采样"],
            ["13", "测试管理平台(FastAPI + Vue 3: 查用例 / 跑用例 / 看报告)",
             "Web 工作台, 不写代码也能跑测试"],
        ],
        col_widths=[1.5, 7.0, 6.5])

    add_para(doc,
        "阶段 13 是为了让不写代码的测试同事也能用起来. 之前要跑测试得敲命令行, "
        "现在浏览器打开 http://localhost:5173 就能查项目、选用例、触发执行、看报告.",
        indent=0.7)

    add_callout(doc,
        "整个平台搭建过程的对话累积留痕在 docs/协作留痕.docx, 各国业务测试留痕在 "
        "docs/业务留痕-<国家>.docx. 这些是跨会话记忆, 我下一次继续开发时会先读这些恢复上下文.",
        label="留痕机制")
    add_page_break(doc)


def build_part4(doc: Document) -> None:
    add_heading(doc, "第四部分  核心设计思路详解", level=1)

    add_heading(doc, "4.1  framework + projects 的分层逻辑", level=2)
    add_para(doc,
        "这是整个平台最重要的设计决策. 想象一下公司里有 5 个业务系统都要做自动化, "
        "如果不分层, 会变成这样:", indent=0.7)

    add_code_block(doc, """分层前(每个项目各自一套):
  oms/   ├─ http_client.py (自己写的)
         ├─ logger.py       (自己写的)
         ├─ assertions.py   (自己写的)
         └─ testcase/        (业务用例)
  wms/   ├─ http_client.py (又写了一遍, 跟 oms 的还不一样)
         ├─ logger.py       (又写一遍)
         └─ ...
  → 重复造轮子, 标准不统一, 升级要改 5 处

分层后(framework 沉淀公共能力):
  framework/
    ├─ clients/http/   (统一 HTTP 客户端)
    ├─ core/           (配置/日志/上下文)
    └─ testing/        (断言/数据驱动)
  projects/
    ├─ oms/testcase/   (只写业务用例, 用 framework 的能力)
    ├─ wms/testcase/   (同上)
    └─ ...
  → 升级 framework 一处, 所有项目受益""")

    add_para(doc,
        "framework 是 uv workspace 的成员, projects 通过 [tool.uv.sources] "
        "framework = { workspace = true } 引用它, 改 framework 立即生效, 不用重新安装.",
        indent=0.7)

    add_heading(doc, "4.2  数据驱动与接口依赖", level=2)

    add_para(doc, "数据驱动: 把测试数据从代码里抽出来放 YAML, 一份代码跑多组数据.",
             bold=True, indent=0.7)
    add_code_block(doc, """# data/login_cases.yaml  (数据)
cases:
  - id: admin_login
    username: admin
    password: admin123
  - id: user_login
    username: user01
    password: user123

# testcase/test_login.py  (代码)
@pytest.mark.parametrize("case", load_cases(yaml_path), ids=case_ids(yaml_path))
async def test_login(http_client, case):
    token = await LoginApi(http_client).login(case["username"], case["password"])
    assert token.access_token != ""
""")
    add_para(doc,
        "好处: 加测试场景只改 YAML 不改代码; 同一份代码能跑 admin 和 user 两种角色.",
        indent=0.7)

    add_para(doc, "接口依赖: 用 DAG(有向无环图)自动按依赖顺序执行.", bold=True, indent=0.7)
    add_para(doc,
        "举例: 创建订单接口依赖\"登录拿 token\"和\"查询商品 ID\", 而查询商品又依赖\"创建商品\". "
        "传统做法是手动写顺序, 容易出错. 我们用拓扑排序自动计算执行顺序, "
        "前一接口的响应字段还能自动提取注入下一接口.", indent=0.7)

    add_heading(doc, "4.3  并发执行与 Allure 报告", level=2)
    add_para(doc, "并发: pytest-xdist 多进程并行跑, -n auto 自动按 CPU 核数分配.",
             bold=True, indent=0.7)
    add_code_block(doc, """# 串行(慢)
uv run pytest                              # 100 个用例要跑 10 分钟

# 并发(快)
uv run pytest -n auto                      # 同样 100 个用例 2 分钟跑完
""")
    add_para(doc,
        "进程级并行的好处是隔离性好, 一个用例崩了不影响其他; 缺点是不能共享内存状态, "
        "但对接口测试没影响.", indent=0.7)

    add_para(doc, "Allure 报告: 自动生成富文本报告, 含步骤/截图/标签/历史趋势.",
             bold=True, indent=0.7)
    add_bullet(doc, "step(步骤): with step('登录') 包裹, 报告里能看到业务流程")
    add_bullet(doc, "attach(附件): HTTP 请求/响应自动作为附件, 失败截图自动附上")
    add_bullet(doc, "labels(标签): 用 @allure.suite / @allure.severity 分类")
    add_bullet(doc, "environment: 测试环境(base_url / browser / channel)写进报告")
    add_bullet(doc, "categories: 失败分类(产品缺陷/测试缺陷/Flaky)")
    add_bullet(doc, "history: 历史趋势, 看通过率随时间变化")

    add_heading(doc, "4.4  WebUI/App 扩展", level=2)
    add_para(doc,
        "Playwright 做 Web 自动化, Appium 做 App 自动化. 两个客户端都封装成统一接口, "
        "业务代码切换 Web/App 只改 fixture, 不改用例.", indent=0.7)

    add_para(doc, "腾讯滑块验证码求解器(亮点):", bold=True, indent=0.7)
    add_para(doc,
        "jmseu(德国 JMS)登录会触发腾讯滑块验证码. 我先用 OpenCV 做了缺口检测 + "
        "人类轨迹模拟自动求解, 但发现两点: (1) 太慢影响测试效率; (2) 容易被封 IP. "
        "后来改成人机协作模式: 操作员在可见浏览器手动完成滑动, 脚本等待 dashboard 出现 "
        "'功能入口' 文本即认为登录成功. 这是工程上的折中, 不追求全自动, 解决问题优先.",
        indent=0.7)

    add_heading(doc, "4.5  AI 失败分析", level=2)
    add_para(doc,
        "测试失败后, 排错很费时间. 我接入了大模型做失败分析: 用例失败时, "
        "LLMAnalyzer 把失败堆栈、日志、最近 HTTP 交互喂给 LLM, 自动给出可能原因和修复建议. "
        "这降低了新人排错的门槛 - 不用读完整代码也能定位问题.", indent=0.7)
    add_para(doc,
        "设计上用 ABC 抽象 FailureAnalyzer, NullAnalyzer 是空实现(不调 LLM 时用), "
        "LLMAnalyzer 是真实实现. 这样接入不同 LLM 只改实现类, 业务代码不感知.",
        indent=0.7)
    add_page_break(doc)


def build_part5(doc: Document) -> None:
    add_heading(doc, "第五部分  一个贯穿示例: 从写用例到看报告", level=1)

    add_para(doc,
        "用一个登录接口测试贯穿整个流程, 让大家看清楚从写代码到看报告的完整链路. "
        "这个示例来自 projects/template, 是新业务系统的模板.", indent=0.7)

    add_heading(doc, "5.1  写一个登录接口测试", level=2)
    add_para(doc, "第一步: 封装登录接口(api/login.py).", bold=True, indent=0.7)
    add_code_block(doc, """# projects/template/api/login.py
from framework.clients.http.auth import Token
from framework.clients.http.client import AsyncHttpClient

class LoginApi:
    \"\"\"登录接口封装.\"\"\"

    def __init__(self, client: AsyncHttpClient) -> None:
        self._client = client

    async def login(self, username: str, password: str) -> Token:
        resp = await self._client.post(
            "/login", json={"username": username, "password": password}
        )
        resp.raise_for_status()
        return Token(access_token=resp.json["access_token"])
""")

    add_para(doc, "第二步: 写测试用例(testcase/test_login.py).", bold=True, indent=0.7)
    add_code_block(doc, """# projects/template/testcase/test_login.py
import pytest
from api.login import LoginApi
from framework.clients.http.client import AsyncHttpClient

@pytest.mark.smoke  # 标记为冒烟测试, CI 每次提交都跑
async def test_login_returns_token(http_client: AsyncHttpClient) -> None:
    token = await LoginApi(http_client).login("demo", "demo123")
    assert token.access_token == "mock-token-xyz"
""")

    add_para(doc, "第三步: 提供 fixture(fixture/clients.py).", bold=True, indent=0.7)
    add_code_block(doc, """# projects/template/fixture/clients.py
import pytest
from framework.clients.http.client import AsyncHttpClient

@pytest.fixture
async def http_client(mock_transport):
    \"\"\"默认用 mock transport, 不连真实后端, CI 也能跑.\"\"\"
    async with AsyncHttpClient(
        settings=HttpSettings(base_url="http://api.example.com"),
        transport=mock_transport,
    ) as client:
        yield client
""")

    add_heading(doc, "5.2  跑起来", level=2)
    add_para(doc, "命令行方式(开发用):", bold=True, indent=0.7)
    add_code_block(doc, """# 进入项目目录
cd projects/template

# 跑全部用例
..\\..\\.venv\\Scripts\\python.exe -m pytest

# 只跑冒烟
..\\..\\.venv\\Scripts\\python.exe -m pytest -m smoke

# 并发跑
..\\..\\.venv\\Scripts\\python.exe -m pytest -n auto

# 详细输出
..\\..\\.venv\\Scripts\\python.exe -m pytest -v
""")

    add_para(doc, "Web 平台方式(非开发同事用):", bold=True, indent=0.7)
    add_para(doc,
        "浏览器打开 http://localhost:5173 → 用例查询 → 选 template 项目 → "
        "勾选要跑的用例 → 点\"执行测试\" → 实时看日志 → 跑完跳报告页面. 全程不写一行代码.",
        indent=0.7)

    add_heading(doc, "5.3  看 Allure 报告", level=2)
    add_para(doc, "测试跑完会自动生成 Allure HTML 报告, 包含:", indent=0.7)
    add_bullet(doc, "Overview: 总览, 通过率/失败率/用例数")
    add_bullet(doc, "Categories: 按失败类型分类(产品缺陷/测试缺陷/Flaky)")
    add_bullet(doc, "Suites: 按测试套件分组")
    add_bullet(doc, "Timeline: 时间轴, 看每个用例耗时")
    add_bullet(doc, "Behaviors: 按业务功能分组(epic/feature/story)")
    add_bullet(doc, "具体用例: 点进去看步骤、附件(HTTP 请求响应、失败截图)")

    add_para(doc, "终端会自动打开本地 HTTP 服务器浏览报告(避免 file:// 跨源 fetch 转圈):",
             bold=True, indent=0.7)
    add_code_block(doc, """# 控制台会输出报告绝对路径和打开命令:
# Report: C:\\work\\...\\allure-report\\index.html
# Server: http://localhost:5274
""")
    add_callout(doc,
        "Allure 报告是 SPA(单页应用), 直接用浏览器 file:// 打开会一直转圈(Chrome 禁止跨源 fetch). "
        "所以我用 allure open 启动本地 HTTP 服务器打开, 这个细节在 conftest.py 的钩子里自动处理了.",
        label="踩坑提示")
    add_page_break(doc)


def build_part6(doc: Document) -> None:
    add_heading(doc, "第六部分  代码管理与协同", level=1)

    add_heading(doc, "6.1  三套 Git 远程: 个人 / 本地 GitLab", level=2)
    add_para(doc,
        "我把代码同时维护在三个地方, 各有不同用途, 完全隔离互不污染:",
        indent=0.7)
    add_table(doc,
        headers=["远程", "地址", "用途", "谁在用"],
        rows=[
            ["gitee", "gitee.com/wsywpp/automation-platform.git", "个人备份与版本管理",
             "我本人"],
            ["github", "github.com/Rogers501/automation-platform.git", "个人备份镜像",
             "我本人"],
            ["gitlab(本地)", "http://10.66.67.26:8929/root/automation-platform.git",
             "部门协同开发", "部门同事"],
        ],
        col_widths=[2.5, 5.5, 4.0, 3.0])

    add_para(doc, "本地 Git 配置:", bold=True, indent=0.7)
    add_code_block(doc, """# 看当前所有远程
git remote -v
# origin   → gitee + github (个人仓库, 多 push URL)
# gitlab   → 本地 GitLab (部门协同)

# 推到个人仓库
git push origin main

# 推到部门本地 GitLab
git push gitlab main

# 两条路径完全独立, 推 gitlab 不会污染 gitee/github
""")

    add_heading(doc, "6.2  分支与提交规范", level=2)
    add_para(doc, "分支策略(轻量级, 适合小团队):", bold=True, indent=0.7)
    add_bullet(doc, "main: 主分支, 一直可跑, 一直可发布")
    add_bullet(doc, "feature/xxx: 新功能分支, 完成后合 main")
    add_bullet(doc, "fix/xxx: bug 修复分支")
    add_bullet(doc, "hotfix/xxx: 紧急修复, 直接合 main")

    add_para(doc, "提交信息规范(Conventional Commits):", bold=True, indent=0.7)
    add_code_block(doc, """<类型>: <简短描述>

类型枚举:
  feat:     新功能
  fix:      bug 修复
  docs:     文档变更
  style:    代码格式(不影响功能)
  refactor: 重构
  test:     增加测试
  chore:    构建/工具链变更

示例:
  feat: 增加 jmseu 德国登录测试
  fix: 修复 Kafka 消费者连接泄漏
  docs: 补充部署与运维文档
""")

    add_heading(doc, "6.3  Pre-commit 自动检查", level=2)
    add_para(doc,
        "提交代码前会自动跑 ruff check + ruff format --check + mypy, 任何一项不通过提交被拒绝. "
        "这保证了进 main 的代码永远是干净的.", indent=0.7)
    add_code_block(doc, """# 安装 pre-commit hook(一次性)
uv run pre-commit install

# 之后每次 git commit 会自动跑:
# - ruff check .          (静态检查)
# - ruff format --check .  (格式检查)
# - mypy                   (类型检查)

# 全部通过才提交成功, 不通过看报错修复后再提交
""")

    add_heading(doc, "6.4  CI/CD 流水线", level=2)
    add_para(doc,
        "代码推到 GitLab 后, 自动触发 CI 流水线. 流水线分四个阶段:",
        indent=0.7)
    add_code_block(doc, """提交代码
   ↓
[lint]   ruff check + ruff format --check + mypy
   ↓ (失败阻断)
[smoke]  pytest -m smoke -n auto         (冒烟, 快速门禁)
   ↓ (失败阻断)
[regression]  pytest -m regression -n auto  (全量回归)
   ↓
[report] allure generate → GitLab Pages
""")

    add_para(doc, "关键设计:", bold=True, indent=0.7)
    add_bullet(doc,
        "单入口: 所有命令集中在 scripts/ci/ci.sh, GitLab CI 和 Jenkins 共用, 改命令只动一处")
    add_bullet(doc, "测试分级: smoke(冒烟, 每次提交跑) / regression(回归, 全量跑)")
    add_bullet(doc, "Allure 报告自动发布到 GitLab Pages, 推完代码过几分钟就能看报告")
    add_bullet(doc, "Jenkins 用 Allure Plugin, agent 上不用装 allure CLI")
    add_page_break(doc)


def build_part7(doc: Document) -> None:
    add_heading(doc, "第七部分  部署与运维", level=1)

    add_heading(doc, "7.1  本地依赖(MySQL/Redis/Kafka)", level=2)
    add_para(doc,
        "测试框架本身需要 MySQL / Redis / Kafka 作为测试目标(测 DB 校验、缓存、消息队列). "
        "全部用 docker-compose 一键拉起:", indent=0.7)
    add_code_block(doc, """# 启动全部依赖(MySQL + Redis + Kafka)
docker compose -f docker/docker-compose.yml up -d

# 只启动 MySQL
docker compose -f docker/docker-compose.yml up -d mysql

# 看状态
docker compose -f docker/docker-compose.yml ps
""")
    add_para(doc,
        "MySQL 数据持久化到项目本地 data/mysql/, Docker Desktop 升级或容器重建都不会丢数据.",
        indent=0.7)

    add_heading(doc, "7.2  测试管理平台一键启动", level=2)
    add_para(doc,
        "测试管理平台(FastAPI 后端 + Vue 前端)封装了一键启动脚本:",
        indent=0.7)
    add_code_block(doc, """# 一键启动(MySQL 容器 + 后端 + 前端)
powershell -ExecutionPolicy Bypass -File scripts/start_all.ps1

# 脚本自动完成:
# 1. 检查并加载 MySQL 离线镜像(避免 Docker Hub 拉不下来)
# 2. 启动 MySQL 容器, 等待健康检查
# 3. 启动 FastAPI 后端(uvicorn --port 8900)
# 4. 启动 Vite 前端(--port 5173)
# 5. 轮询 /api/health 确认后端就绪

# 访问入口:
# 前端工作台: http://localhost:5173
# 后端 API 文档: http://localhost:8900/docs
# 健康检查: http://localhost:8900/api/health
""")

    add_heading(doc, "7.3  本地 GitLab 容器", level=2)
    add_para(doc,
        "为了部门协同开发, 不污染我个人 gitee/github 仓库, 我在本地 Docker 起了一个 GitLab CE:",
        indent=0.7)
    add_table(doc,
        headers=["项", "值"],
        rows=[
            ["Web UI", "http://10.66.67.26:8929"],
            ["HTTP 克隆", "http://10.66.67.26:8929/root/automation-platform.git"],
            ["SSH 克隆", "ssh://git@10.66.67.26:2224/root/automation-platform.git"],
            ["管理员账号", "root(首次密码见容器内 /etc/gitlab/initial_root_password)"],
            ["同事共享账号", "team / AutoDev2026!Xq (Maintainer 权限)"],
            ["docker-compose", "docker/gitlab/docker-compose.yml"],
            ["数据持久化", "C:/docker-data/git-server/{config,data,logs}"],
            ["镜像备份", "C:/docker-data/git-server/images/gitlab-ce-latest.tar (1.4GB)"],
        ],
        col_widths=[4.0, 11.0])

    add_callout(doc,
        "GitLab 镜像约 2.5GB, Docker Hub 国内拉不下来, 我通过 docker.1ms.run 镜像源拉到本地后, "
        "又 docker save 导出成 tar 文件, 防止 Docker GC 删除. 同事换机器也能 docker load 恢复.",
        label="镜像管理")

    add_heading(doc, "7.4  数据备份与镜像管理", level=2)
    add_para(doc, "MySQL 备份:", bold=True, indent=0.7)
    add_code_block(doc, """# 备份(导出 SQL)
powershell -ExecutionPolicy Bypass -File scripts/backup_mysql.ps1
# 产物: data/backups/automation-<timestamp>.sql

# 恢复
powershell -ExecutionPolicy Bypass -File scripts/restore_mysql.ps1 \\
    -BackupFile data\\backups\\automation-20260823-100000.sql
""")

    add_para(doc, "MySQL 镜像导出(防止 Docker 重装丢镜像):", bold=True, indent=0.7)
    add_code_block(doc, """# 导出
powershell -ExecutionPolicy Bypass -File scripts/save_mysql_image.ps1
# 产物: data/docker-images/mysql-8.4.tar

# 恢复
docker load -i data\\docker-images\\mysql-8.4.tar
""")

    add_para(doc, "建议备份策略:", bold=True, indent=0.7)
    add_table(doc,
        headers=["频率", "范围", "保留期"],
        rows=[
            ["每日", "MySQL 全量 dump", "7 天"],
            ["每周", "整个 data/mysql/ 目录", "4 周"],
            ["重大变更前", "手动备份", "永久"],
        ],
        col_widths=[3.0, 7.0, 4.0])
    add_page_break(doc)


def build_part8(doc: Document) -> None:
    add_heading(doc, "第八部分  同事参与指南", level=1)

    add_heading(doc, "8.1  新人 30 分钟上手", level=2)
    add_para(doc, "按这个顺序走, 半小时内能跑通第一个用例:", indent=0.7)
    add_para(doc, "第一步: 装环境(10 分钟).", bold=True, indent=0.7)
    add_code_block(doc, """# 1. 装 Python 3.12+ (官网下载)
# 2. 装 uv (Python 包管理器)
powershell -ExecutionPolicy ByPass -c "irm https://astral.sh/uv/install.ps1 | iex"

# 3. 装 Git (官网下载)
# 4. 装 Docker Desktop (官网下载, 跑测试依赖用)
# 5. 装 PyCharm Community (官网下载, 推荐) 或 VS Code
""")

    add_para(doc, "第二步: 拉代码(5 分钟).", bold=True, indent=0.7)
    add_code_block(doc, """# 找我要 team 账号密码(team / AutoDev2026!Xq)
git clone http://10.66.67.26:8929/root/automation-platform.git
cd automation-platform

# 装 Python 依赖(创建 .venv + 装 framework + 全部依赖)
uv sync
""")

    add_para(doc, "第三步: PyCharm 配置(5 分钟).", bold=True, indent=0.7)
    add_bullet(doc, "File → Open → 选 automation-platform 目录")
    add_bullet(doc, "Settings → Project → Python Interpreter → Existing")
    add_bullet(doc, "选 C:\\work\\...\\automation-platform\\.venv\\Scripts\\python.exe")
    add_bullet(doc, "Settings → Tools → Python Integrated Tools → Default test runner: pytest")

    add_para(doc, "第四步: 跑第一个测试(10 分钟).", bold=True, indent=0.7)
    add_code_block(doc, """# 跑 framework 自带测试(验证环境)
uv run pytest -m smoke

# 跑 template 项目的示例测试(mock 模式, 不需要真实环境)
cd projects/template
..\\..\\.venv\\Scripts\\python.exe -m pytest
""")
    add_callout(doc,
        "第一个测试跑通后, 接下来读 docs/快速上手.md(30 分钟), "
        "再看 docs/用户手册.md 完整了解. 不懂的技术名词查本文档附录 A.",
        label="下一步")

    add_heading(doc, "8.2  跑测试的三种姿势", level=2)
    add_table(doc,
        headers=["姿势", "适合谁", "命令/操作", "优缺点"],
        rows=[
            ["PyCharm 点绿按钮", "开发/测试开发",
             "在测试函数旁点绿色三角形", "可断点调试, 单个用例快"],
            ["命令行 pytest", "开发/测试开发",
             "uv run pytest -m smoke -n auto", "可批量, 可并发, 灵活"],
            ["Web 平台", "所有人(含非开发)",
             "http://localhost:5173 选项目→选用例→执行", "界面友好, 实时日志, 看报告方便"],
        ],
        col_widths=[3.0, 2.5, 5.5, 4.0])

    add_heading(doc, "8.3  查看报告", level=2)
    add_para(doc, "测试跑完, Allure 报告自动生成, 三种看法:", indent=0.7)
    add_bullet(doc,
        "命令行: 测试跑完控制台会输出报告路径和 allure open 命令, 复制执行即可")
    add_bullet(doc,
        "PyCharm: 跑完测试, 控制台底部有报告绝对路径(可点击)")
    add_bullet(doc,
        "Web 平台: 报告中心页面, 选择项目直接看通过率/失败明细/历史趋势")

    add_heading(doc, "8.4  提交代码流程", level=2)
    add_code_block(doc, """# 1. 拉最新代码
git pull gitlab main

# 2. 创建分支(不要直接在 main 上改)
git checkout -b feature/add-jmseu-login-test

# 3. 写代码 + 写测试
# (在 projects/jmseu/testcase/ 下加测试文件)

# 4. 跑测试验证
uv run pytest projects/jmseu/ -v

# 5. 跑静态检查
uv run ruff check . && uv run ruff format .
uv run mypy

# 6. 提交(pre-commit 会自动跑检查)
git add .
git commit -m "feat: 增加 jmseu 德国登录测试"

# 7. 推送到本地 GitLab
git push gitlab feature/add-jmseu-login-test

# 8. 在 GitLab Web UI 创建 Merge Request, 通知 reviewer

# 9. 合并到 main 后, CI 自动跑全量测试, 报告发布到 Pages
""")

    add_callout(doc,
        "提 MR 前务必: (1) 跑过本地测试; (2) ruff/mypy 全过; (3) 提交信息符合规范. "
        "MR 描述写清楚改了什么、为什么改、怎么验证. reviewer 在 GitLab 上 review 代码后合并.",
        label="协作规范")
    add_page_break(doc)


def build_part9(doc: Document) -> None:
    add_heading(doc, "第九部分  踩过的坎与决策理由", level=1)

    add_para(doc,
        "搭建过程中踩了不少坑, 这里挑几个有代表性的, 希望大家以后避坑.", indent=0.7)

    add_para(doc, "坎 1: Docker Hub 国内拉不下来", bold=True, indent=0.7)
    add_para(doc,
        "现象: docker pull gitlab/gitlab-ce:latest 一直 EOF, 拉到 1.4GB 断流. "
        "解决: 试了多个国内镜像源(dockerproxy / daocloud / 1ms.run), 最终 1ms.run 拉成功. "
        "又 docker save 导出 tar 备份, 防止下次又拉不下来.", indent=0.7)

    add_para(doc, "坎 2: Python 3.12 + setuptools 兼容性", bold=True, indent=0.7)
    add_para(doc,
        "现象: PyCharm 测试运行器报 pkg_resources 找不到. "
        "原因: setuptools v83 移除了 pkg_resources, 但 PyCharm 还在用. "
        "解决: 在 pyproject.toml dev 依赖里固定 setuptools<81.", indent=0.7)

    add_para(doc, "坎 3: Allure 报告 file:// 打开一直转圈", bold=True, indent=0.7)
    add_para(doc,
        "现象: 直接双击 allure-report/index.html, 报告加载不出来. "
        "原因: Allure 是 SPA, 通过 fetch 加载 data/*.json, Chrome 禁止 file:// 跨源 fetch. "
        "解决: conftest.py 钩子里用 allure open 启动本地 HTTP 服务器打开.", indent=0.7)

    add_para(doc, "坎 4: pytest-xdist 进程隔离导致 fixture 共享失败", bold=True, indent=0.7)
    add_para(doc,
        "现象: 串行跑没问题, 并发跑 token 丢失. "
        "原因: xdist 是多进程, 主进程的内存状态不共享给 worker. "
        "解决: token 改用文件持久化, worker 进程从文件读; 或者用 pytest-xdist 的 tmp_path_factory 同步.",
        indent=0.7)

    add_para(doc, "坎 5: 腾讯滑块验证码自动破解被封 IP", bold=True, indent=0.7)
    add_para(doc,
        "现象: 用 OpenCV 自动求解验证码, 跑几次就被封 IP. "
        "解决: 改成人机协作模式, 操作员手动滑, 脚本等待成功标志. "
        "教训: 不要为了全自动牺牲可用性, 工程上解决问题优先.", indent=0.7)

    add_para(doc, "坎 6: GitLab daemon.json 镜像加速被还原", bold=True, indent=0.7)
    add_para(doc,
        "现象: 改了 ~/.docker/daemon.json 加 registry-mirrors, 重启 Docker Desktop 后被还原. "
        "原因: Docker Desktop 用内部引擎, 不直接读这个文件. "
        "解决: 用镜像源前缀直接拉(docker.1ms.run/gitlab/gitlab-ce:latest), 再 retag.",
        indent=0.7)

    add_para(doc, "关键决策回顾:", bold=True, indent=0.7)
    add_table(doc,
        headers=["决策点", "选择", "理由"],
        rows=[
            ["依赖管理", "uv (非 pip/poetry)", "快 10-100 倍, workspace 原生支持"],
            ["HTTP 客户端", "httpx (非 requests)", "原生 async, 自带 MockTransport"],
            ["Web 自动化", "Playwright (非 Selenium)", "原生 async, 自动等待, trace viewer"],
            ["报告", "Allure (非 HTML report)", "富报告, 历史趋势, 行业标准"],
            ["分层", "framework + projects", "公共能力沉淀, 业务独立维护"],
            ["CI 入口", "scripts/ci/ci.sh 单入口", "GitLab/Jenkins 共用, 改一处"],
            ["平台后端", "FastAPI (非 Flask/Django)", "原生 async, 自动 OpenAPI 文档"],
            ["Git 远程", "个人 + 本地 GitLab 隔离", "部门协同不污染个人仓库"],
        ],
        col_widths=[3.0, 4.0, 8.0])
    add_page_break(doc)


def build_part10(doc: Document) -> None:
    add_heading(doc, "第十部分  后续演进与展望", level=1)

    add_para(doc,
        "当前 14 个阶段全部完成, 原始目标全部实现. 但还有几个方向值得继续演进:",
        indent=0.7)

    add_heading(doc, "10.1  测试管理平台演进", level=2)
    add_table(doc,
        headers=["演进项", "说明", "优先级"],
        rows=[
            ["定时调度", "cron 触发每日回归, 不用手动点", "高"],
            ["远程执行节点", "支持多台机器分发执行, 跑大用例集不堵", "中"],
            ["用户与权限", "区分只读/执行/管理角色", "中"],
            ["报告历史趋势", "ECharts 展示多版本通过率曲线", "中"],
            ["测试数据管理 UI", "目前 YAML 手改, 后续做界面化", "低"],
        ],
        col_widths=[3.0, 7.0, 3.0])

    add_heading(doc, "10.2  框架能力演进", level=2)
    add_bullet(doc, "AI 用例生成: 现在只有失败分析, 后续用 LLM 根据 OpenAPI 自动生成用例骨架")
    add_bullet(doc, "性能测试集成: loadtest 已有 Locust, 后续压测结果也接入平台报告")
    add_bullet(doc, "Mock 服务: 目前用 httpx MockTransport, 后续可起独立 mock 服务做联调")
    add_bullet(doc, "测试覆盖率统计: 接入 coverage.py, 看哪些代码路径没覆盖到")
    add_bullet(doc, "更多的 App 测试: 当前 App 集成了 Appium, 但实际用例还少, 后续补充")

    add_heading(doc, "10.3  协同与流程演进", level=2)
    add_bullet(doc, "本地 GitLab 接 CI Runner: 当前只做代码托管, 后续接 runner 跑流水线")
    add_bullet(doc, "定时全量回归: 每晚自动跑全量, 第二天上班看报告")
    add_bullet(doc, "缺陷自动登记: 测试失败自动在 Jira/Linear 建缺陷单")
    add_bullet(doc, "通知集成: 失败自动推送到企业微信/钉钉群")

    add_para(doc, "", )
    add_callout(doc,
        "平台不是终点, 是起点. 真正的价值在于: 让测试同事从手工重复劳动中解放出来, "
        "把时间花在更有价值的探索性测试和业务理解上. 这是我搭建这个平台的初心.",
        label="愿景")
    add_page_break(doc)


def build_appendix_a(doc: Document) -> None:
    add_heading(doc, "附录 A  技术名词解释", level=1)
    add_para(doc,
        "本文档涉及的技术名词集中解释, 方便非开发背景同事查阅.", indent=0.7)

    add_table(doc,
        headers=["名词", "解释"],
        rows=[
            ["Python", "一种通用编程语言, 语法简洁, 适合自动化测试. 我们用 3.12 版本"],
            ["async/await", "异步编程语法, 让程序在等待 IO 时不阻塞, 提升并发性能"],
            ["pytest", "Python 测试框架, 比自带的 unittest 更灵活, 生态丰富"],
            ["fixture", "pytest 的概念, 测试前置准备(如创建 HTTP 客户端), 测试完自动清理"],
            ["uv", "Astral 公司用 Rust 写的 Python 包管理器, 比 pip 快 10-100 倍"],
            ["workspace", "uv 的功能, 让多个 Python 项目共享一个虚拟环境"],
            ["httpx", "Python 的 HTTP 客户端库, 原生支持 async, 类似 requests 但更现代"],
            ["pydantic", "Python 数据校验库, 用类型注解定义数据结构, 自动校验"],
            ["loguru", "Python 日志库, 零配置, 比自带 logging 简单"],
            ["SQLAlchemy", "Python 的 ORM, 把 SQL 操作封装成 Python 对象"],
            ["Playwright", "微软的浏览器自动化工具, 类似 Selenium 但更现代"],
            ["Appium", "App 自动化工具, 支持 iOS/Android"],
            ["Kafka/RabbitMQ/RocketMQ", "三种主流消息队列, 用于系统间异步通信"],
            ["Redis", "内存数据库, 常用作缓存"],
            ["Docker", "容器化工具, 把应用和依赖打包成可移植的容器"],
            ["docker-compose", "Docker 的多容器编排工具, 一个命令启动多个服务"],
            ["Git", "分布式版本控制系统, 管理代码版本"],
            ["GitLab", "Git 仓库管理平台, 类似 GitHub 但可私有部署"],
            ["GitLab CI", "GitLab 内置的持续集成功能, 推代码自动跑流水线"],
            ["Jenkins", "开源的持续集成服务器, 老牌但功能强大"],
            ["FastAPI", "Python 的现代 Web 框架, 原生 async, 自动生成 API 文档"],
            ["Vue 3", "前端框架, 渐进式, 国内主流"],
            ["Element Plus", "基于 Vue 3 的 UI 组件库, 表格表单弹窗齐全"],
            ["Vite", "前端构建工具, 比 webpack 快, 开发体验好"],
            ["Allure", "测试报告框架, 富文本报告, 历史趋势, 行业标准"],
            ["pytest-xdist", "pytest 插件, 多进程并行跑测试"],
            ["ruff", "Rust 写的 Python 静态检查+格式化工具, 替代 flake8+isort+black"],
            ["mypy", "Python 类型检查工具, 在运行前发现类型错误"],
            ["pre-commit", "Git hook 管理工具, 提交前自动跑检查"],
            ["YAML", "数据序列化格式, 类似 JSON 但更易读, 用缩进表示层级"],
            ["JSON", "数据交换格式, 浏览器和服务器通信的标准格式"],
            ["OpenAPI", "API 描述规范, 自动生成接口文档和客户端代码"],
            ["WebSocket", "全双工通信协议, 服务器可主动推送数据到浏览器"],
            ["ORM", "Object Relational Mapping, 把数据库表映射成对象"],
            ["ABC", "Abstract Base Class, Python 的抽象基类, 定义接口规范"],
            ["DAG", "Directed Acyclic Graph, 有向无环图, 用于接口依赖排序"],
            ["CI/CD", "Continuous Integration / Continuous Deployment, 持续集成/部署"],
            ["Merge Request (MR)", "GitLab 的合并请求, 类似 GitHub 的 Pull Request"],
            ["SPA", "Single Page Application, 单页应用, 全程不刷新页面"],
            ["Mock", "模拟, 用假数据/假服务代替真实后端, 方便测试"],
            ["Token", "登录凭证, 登录成功后服务器返回, 后续请求带上证明身份"],
            ["BearerAuth", "HTTP 认证方式, 请求头带 Authorization: Bearer <token>"],
            ["TLS/HTTPS", "传输层加密, HTTPS = HTTP + TLS, 防止数据被中间人窃听"],
            ["OpenCV", "计算机视觉库, 用于图像识别(如验证码缺口检测)"],
            ["Locust", "Python 的压测工具, 用代码定义压测场景"],
            ["SLA", "Service Level Agreement, 服务水平协议, 如 99% 请求 200ms 内响应"],
        ],
        col_widths=[3.5, 12.0])
    add_page_break(doc)


def build_appendix_b(doc: Document) -> None:
    add_heading(doc, "附录 B  常用命令速查", level=1)

    add_heading(doc, "环境与依赖", level=2)
    add_code_block(doc, """uv sync                                   # 装 Python 依赖
uv add <package>                          # 加新依赖
uv export --frozen -o requirements.txt    # 导出 requirements.txt
""")

    add_heading(doc, "跑测试", level=2)
    add_code_block(doc, """uv run pytest                             # 跑全部用例
uv run pytest -m smoke                    # 只跑冒烟
uv run pytest -m regression               # 只跑回归
uv run pytest -n auto                     # 并发跑
uv run pytest projects/jmseu/             # 指定项目
uv run pytest -k test_login               # 按名过滤
uv run pytest -v -s                       # 详细输出+显示 print
""")

    add_heading(doc, "代码质量", level=2)
    add_code_block(doc, """uv run ruff check .                       # 静态检查
uv run ruff format .                      # 格式化
uv run ruff check --fix . && uv run ruff format .   # 一键修复
uv run mypy                               # 类型检查
""")

    add_heading(doc, "Docker 与依赖服务", level=2)
    add_code_block(doc, """docker compose -f docker/docker-compose.yml up -d      # 启动测试依赖
docker compose -f docker/docker-compose.yml ps        # 看状态
docker compose -f docker/docker-compose.yml down      # 停止

cd docker/gitlab && docker compose up -d              # 启动本地 GitLab
docker ps --filter "name=gitlab"                       # 看 GitLab 状态
""")

    add_heading(doc, "测试管理平台", level=2)
    add_code_block(doc, """# 一键启动
powershell -ExecutionPolicy Bypass -File scripts/start_all.ps1

# 分进程启动
powershell -ExecutionPolicy Bypass -File scripts/start_backend.ps1
powershell -ExecutionPolicy Bypass -File scripts/start_frontend.ps1

# 健康检查
curl http://localhost:8900/api/health
""")

    add_heading(doc, "备份与镜像", level=2)
    add_code_block(doc, """# MySQL 备份/恢复
powershell -ExecutionPolicy Bypass -File scripts/backup_mysql.ps1
powershell -ExecutionPolicy Bypass -File scripts/restore_mysql.ps1 -BackupFile <file>

# MySQL 镜像导出/导入
powershell -ExecutionPolicy Bypass -File scripts/save_mysql_image.ps1
docker load -i data\\docker-images\\mysql-8.4.tar

# GitLab 镜像导入
docker load -i C:/docker-data/git-server/images/gitlab-ce-latest.tar
""")

    add_heading(doc, "Git 操作", level=2)
    add_code_block(doc, """git remote -v                             # 看所有远程
git push origin main                      # 推个人 gitee/github
git push gitlab main                      # 推部门本地 GitLab
git pull gitlab main                      # 拉最新代码
git checkout -b feature/xxx               # 创建新分支
git add . && git commit -m "feat: xxx"    # 提交
""")

    add_heading(doc, "Allure 报告", level=2)
    add_code_block(doc, """# 生成 HTML 报告(测试跑完自动生成, 一般不用手动)
sh scripts/ci/ci.sh report

# 本地服务器打开(避免 file:// 转圈)
allure open allure-report
allure open --port 5274 allure-report
""")

    add_heading(doc, "CI 复现", level=2)
    add_code_block(doc, """sh scripts/ci/ci.sh install              # 装依赖
sh scripts/ci/ci.sh lint                  # 跑 lint
sh scripts/ci/ci.sh smoke                # 跑冒烟
sh scripts/ci/ci.sh regression           # 跑回归
sh scripts/ci/ci.sh install-allure       # 装 allure CLI
sh scripts/ci/ci.sh report               # 生成报告
""")

    # 结尾
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— 文档结束 ——")
    set_run_font(run, size=12, color=RGBColor(0x80, 0x80, 0x80))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("如有疑问, 欢迎找我当面交流")
    set_run_font(run, size=11, color=RGBColor(0x80, 0x80, 0x80))


# ---------- 主流程 ----------


def main() -> None:
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 构建文档
    build_cover(doc)
    add_page_break(doc)
    build_toc(doc)
    build_part1(doc)
    build_part2(doc)
    build_part3(doc)
    build_part4(doc)
    build_part5(doc)
    build_part6(doc)
    build_part7(doc)
    build_part8(doc)
    build_part9(doc)
    build_part10(doc)
    build_appendix_a(doc)
    build_appendix_b(doc)

    # 保存
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"文档已生成: {OUTPUT}")
    print(f"文件大小: {OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
